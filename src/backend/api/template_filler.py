import os
import re

import docx

from src.backend.api.base import BaseAPI
from src.backend.api.swal import SwalAPI


class TemplateFiller(BaseAPI):
    """API for working with Word templates"""

    def __init__(self, window, swal):
        super().__init__(window)
        self.swal: SwalAPI = swal
        self.placeholder_regex = r'%\{(.+?)\}%'
        self.output_file = 'output.docx'

    def _open_file(self, file_path):
        try:
            return docx.Document(file_path)
        except Exception as e:
            self.swal.error(f'Erro ao abrir o arquivo: {e}')

    def placeholder(self, placeholder):
        return f'%{{{placeholder}}}%'

    def placeholders(self, file_path):
        doc = self._open_file(file_path)

        placeholders = [
            placeholder
            for paragraph in doc.paragraphs
            for placeholder in re.findall(
                self.placeholder_regex, paragraph.text, re.DOTALL
            )
        ]

        return placeholders

    def fill(self, file_path, data):
        # Check if all values are not empty
        if any(not v for v in data.values()):
            return self.swal.error('Todos os campos devem ser preenchidos!')

        doc = self._open_file(file_path)

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for name, value in data.items():
                    run.text = run.text.replace(self.placeholder(name), value)

        doc.save(self.output_file)
        os.startfile(self.output_file)
